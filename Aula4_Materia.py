# Bibliotecas
import os # Importante para lidar com pastas e ficheiros
from docx import Document # Biblioteca para a leitura e escrita de documentos docx (pip install python-docx)
from PyPDF2 import PdfReader # Biblioteca para ler texto simples em PDF (pip install PyPDF2)
import pdfplumber # Biblioteca para ler texto estruturado e tabelas em PDF (pip install pdfplumber)
from pdf2image import convert_from_path #OCR para a PDFs resultantes de scans
import pytesseract # OCR engine
from fpdf import FPDF # 
import csv # Utilizado para ficheiros csv
from openpyxl import load_workbook # Para ficheiros xlsx (necessário pip install openpyxl)
from openpyxl import Workbook # Para ficheiros xlsx
import xml.etree.ElementTree as ET # Para ficheiros xml
import json # Biblioteca para exercicios que requerem json



# Ficheiros

# Listar todos os ficheiros de uma pasta

def listar_ficheiros(directoria):
    ficheiros = []

    for nome in os.listdir(directoria):
        caminho = os.path.join(directoria, nome)
        if os.path.isfile(caminho):
            ficheiros.append(caminho)

    return ficheiros

	
# Verificar a natureza de ficheiros (.pdf, .docx, .txt):

def extensao_ficheiro(directoria):
    _, ext = os.path.splitext(directoria)
    return ext.lower()

# Leitura de ficheiros:

# Ficheiros txt

def get_texto_txt(path):
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

	
def get_linhastexto_txt(path):
    with open(path, "r", encoding="utf-8") as f:
        return f.readlines()
	
def get_line_text(path, i):
	f = open(path)
	
	return f.readlines()[i]
	

# Se quiser ler linha a linha

def percorrer_linha_a_linha(path):

	linhas = []

	with open(path) as f1:

		lines = f1.readlines()
	
		for line in lines:
			#Código sobre o que fazer com a linha aqui
			linhas.append(line)
			
		return linhas
		
# Escrita de documentos txt

def write_text_txt(path, content, type):

	with open(path, type, encoding="utf-8") as f:
		f.write(content)
		
def write_lines_txt(path, content, type):
	with open(path, type, encoding="utf-8") as f:
		for item in content:
			f.write(item + '\n')
	
		
		

# Leitura de ficheiros docx (Necessário instalar a biblioteca pip install python-docx)

# Ler texto de um ficheiro .docx (Paragrafos num array)

def get_textoarray_docx(path):

	doc = Document(path)

	texto = []

	for paragrafo in doc.paragraphs:
		texto.append(paragrafo.text)

	return texto
	
def get_linhas_docx(path):
    doc = Document(path)
    linhas = []

    for p in doc.paragraphs:
        for linha in p.text.split("\n"):
            linhas.append(linha + '\n')

    return linhas

# Ler o texto todo completo numa variável

def get_textostring_docx(path):
	
	doc = Document(path)

	conteudo = "\n".join(p.text for p in doc.paragraphs)

	return conteudo

# Ler conteudo de tabelas dentro de um docx

def get_tabelas_docx(path):
	doc = Document(path)
	return doc.tables
	
def get_table_docx(path,i):
	doc = Document(path)
	return doc.tables[i]
	
def get_rows_of_table(path, i):
	doc = Document(path)
	return doc.tables[i].rows

def get_row_of_table(path, i, j):
	doc = Document(path)
	return doc.tables[i].rows[j]
	
def get_cells_of_table(path,i):
	doc = Document(path)
	
	celulas = []
	
	for linha in doc.tables[i].rows:
		for celula in linha.cells:
			celulas.append(celula)
			
	return celulas

def get_all_cells_of_documet(path):	

	celulas = []

	doc = Document(path)
	for tabela in doc.tables:
		for linha in tabela.rows:
			for celula in linha.cells:
				print(celula.text)
				
	return celulas
			
# Escrever em documentos docx

# Criar um novo documento e acrescentar conteudo / Substiuir conteudo (overwrite)

def create_document_simpletext_docx(path, paragraphs):
	doc = Document()  # cria um documento novo
	
	for p in paragraphs:
		doc.add_paragraph(p)
		
	doc.save(path)

# Abrir um ficheiro existente e acrescentar conteúdo

def add_paragraphs_simpletext_docx(path, paragarphs):

	doc = Document(path)

	for p in paragraphs:
		doc.add_paragraph(p)

	doc.save(path)
	
# Ler Documentos PDF (Necessário as bibliotecas acima)

# Ler todo o texto de pdf (texto simples)

def get_simpletext_content_pdf(path):

	reader = PdfReader(path)
	result = ""
	
	for page in reader.pages:
		result += page.extract_text() + "\n"
		
	return result

def get_lines_pdf(path):
    reader = PdfReader(path)
    linhas = []

    for page in reader.pages:
        texto = page.extract_text()
        if texto:
            for linha in texto.splitlines():
                linhas.append(linha + '\n')

    return linhas

	
# Ler texto de uma página em especifico

def get_simpletext_content_page_pdf(path,i):
	
	reader = PdfReader(path)
	
	return reader.pages[i].extract_text()
	
# Ler texto estruturado (documento todo)

def get_structuredtext_pdf(path):

	texto = ""

	with pdfplumber.open(path) as pdf:
		for page in pdf.pages:
			texto += page.extract_text() + "\n"
			
	return texto

# Ler texto estruturado (página especifica)

def get_structuredtext__page_pdf(path,i):
	
	with pdfplumber.open(path) as pdf:
		return pdf.pages[i].extract_text()
		
	
# Extração de uma tabela (página especifica)

def get_extractedtable_page_pdf(path,i):

	with pdfplumber.open("C:/Users/djpmg/Desktop/alunos.pdf") as pdf:
	
		page_i = pdf.pages[i]
		tabela = page_i.extract_table()
		
		return tabela
		
def get_text_content_of_digitalized_pdfs(path):

	pages = convert_from_path(path)
	
	result = ""
	for page in pages:
		result += pytesseract.image_to_string(page)

	return result
	
# Create and write pdf file

# Classe auxiliar para mudar de pagina sempre que excede o conteudo
class AutoPDF(FPDF):
    def write_line(self, text, line_height=10):
        if self.get_y() + line_height > self.page_break_trigger:
            self.add_page()
        self.cell(0, line_height, txt=text, ln=True)
		
def create_document_simpletext_pdf(path, content, fonte, tamanho_letra, tamanho_linha):
	
	pdf = AutoPDF()
	pdf.add_page()
	pdf.set_font(fonte, size=tamanho_letra)

	for line in content:
		pdf.write_line(line)

	pdf.output(path)
	

# Ficheiros CSV

# leitura das linhas do ficheiros .csv

def ler_linhas_csv(path):

	linhas = []

	with open("dados.csv", "r", newline="", encoding="utf-8") as ficheiro:
		leitor = csv.reader(ficheiro)
		for linha in leitor:
			linhas.append(linha)
	
	return linhas
	
# fazer update de conteúdo de ficheiros csv (Ler e reesercrever o ficheiro, basicamente)
def update_ficheiro_csv(path, content):
	
	linhas = []
	
	with open("dados.csv", "r", newline="", encoding="utf-8") as ficheiro:
		leitor = csv.reader(ficheiro)
		for linha in leitor:
			linhas.append(linha)
			
	# Modificar as linhas com o conteúdo novo
	
	for l in content:
		linhas.append(l)
		
	with open("dados.csv", "w", newline="", encoding="utf-8") as ficheiro:
		escritor = csv.writer(ficheiro)
		escritor.writerows(linhas)
		
# Escrever num ficheiro (linha a linha). Se não existir, cria

def escrever_ficheiro_csv(path, content):

	with open("novo.csv", "w", newline="", encoding="utf-8") as ficheiro:
		escritor = csv.writer(ficheiro)
		
		for linha in content:
			escritor.writerow(linha)
			
			
# Ficheiros XLSX (Excel moderno posterior a 2007)

# Ler tabelas excel para lista de dicionários

def ler_tabela_xlsx(path, nome_folha):
	
	wb = load_workbook(path)
    ws = wb[nome_folha]

    dados = []
    cabecalho = []

    # Ler cabeçalho (primeira linha)
    for celula in ws[1]:
        cabecalho.append(celula.value)

    # Ler restantes linhas
    for linha in ws.iter_rows(min_row=2, values_only=True):
        dicionario = dict(zip(cabecalho, linha))
        dados.append(dicionario)

    return dados
	
# Escrever ficheiro xlsx

def escrever_ficheiro_xlsx(path, content, nome_folha):
	
	# Se o ficheiro não existir, criar um novo
    try:
        wb = load_workbook(path)
    except FileNotFoundError:
        wb = Workbook()

    # Criar ou selecionar a folha
    if nome_folha in wb.sheetnames:
        ws = wb[nome_folha]
    else:
        ws = wb.create_sheet(nome_folha)

    # Limpar conteúdo anterior
    ws.delete_rows(1, ws.max_row)

    # Obter cabeçalho a partir das chaves do primeiro dicionário
    cabecalho = list(content[0].keys())
    ws.append(cabecalho)

    # Escrever cada linha
    for linha_dict in content:
        linha = [linha_dict.get(col, "") for col in cabecalho]
        ws.append(linha)

    # Guardar alterações
    wb.save(path)
	
	
# Ficheiros XML

# Ler ficheiros xml (simples) para dicionário

def ler_ficheiros_xml_simples(path):
    tree = ET.parse(path)
    root = tree.getroot()

    lista = []

    # IMPORTANTE: tens de definir o nome da tag dos itens
    tag_item = root[0].tag  # por exemplo, "aluno"

    for elem in root.findall(tag_item):
        d = {}
        for child in elem:
            d[child.tag] = child.text
        lista.append(d)

    return lista
	
	
# Escrever ficheiros XML (simples):

def escrever_ficheiro_xml(path,content, nome_elemento_raiz, nome_items):

	# Criar elemento raiz
    root = ET.Element(nome_elemento_raiz)

    # Criar um elemento <item> para cada dicionário
    for item_dict in content:
        item_elem = ET.SubElement(root, nome_items)

        # Criar subtags para cada chave do dicionário
        for chave, valor in item_dict.items():
            sub = ET.SubElement(item_elem, chave)
            sub.text = str(valor)

    # Criar a árvore XML e guardar no ficheiro
    tree = ET.ElementTree(root)
    tree.write(path, encoding="utf-8", xml_declaration=True)
	
	
# Ficheiros JSON (Simples)

# Leitura de ficheiros

def ler_ficheiro_json_simples(path):

	with open(path, "r", encoding="utf-8") as f:
		dados = json.load(f)

		# Se o JSON for um objeto com listas internas, tenta extrair a primeira lista
		if isinstance(dados, dict):
			# procurar a primeira lista dentro do dicionário
			for valor in dados.values():
				if isinstance(valor, list):
					return valor
			# se não encontrar listas, devolve o próprio dicionário dentro de uma lista
			return [dados]

		# Se já for uma lista, devolve diretamente
		if isinstance(dados, list):
			return dados

		# Caso raro: outro tipo: embrulhar numa lista
		return [dados]

# Escrever ficheitos

def escrever_ficheiro_json_simples(path, content):

	# Se for lista
    if isinstance(content, list):
        if len(content) == 1:
            dados = content[0]   # apenas um elemento
        else:
            dados = content      # vários elementos
    else:
        # Se já for um dicionário grava como objeto
        dados = content

    # Escrever no ficheiro
    with open(path, "w", encoding="utf-8") as f:
        json.dump(dados, f, ensure_ascii=False, indent=4))